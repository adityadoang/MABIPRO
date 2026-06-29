import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color_hex="CCCCCC"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def set_dashed_borders(table, color_hex="999999"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="dashed" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:bottom w:val="dashed" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:left w:val="dashed" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:right w:val="dashed" w:sz="6" w:space="0" w:color="{color_hex}"/>
            <w:insideH w:val="none"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def add_header_banner(doc, number, title):
    # Table with 1 row, 2 cells
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Set widths: 0.5 inches and 6.0 inches (total 6.5)
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(6.0)
    
    # Format Cell 0 (Green/Teal banner)
    cell_0 = table.cell(0, 0)
    set_cell_shading(cell_0, "0A8481")
    set_cell_margins(cell_0, top=140, bottom=140, left=100, right=100)
    p0 = cell_0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run0 = p0.add_run(number)
    run0.bold = True
    run0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run0.font.size = Pt(14)
    run0.font.name = 'Arial'
    
    # Format Cell 1 (Dark Navy banner)
    cell_1 = table.cell(0, 1)
    set_cell_shading(cell_1, "111E38")
    set_cell_margins(cell_1, top=140, bottom=140, left=150, right=150)
    p1 = cell_1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run1 = p1.add_run(f"  {title.upper()}")
    run1.bold = True
    run1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run1.font.size = Pt(12)
    run1.font.name = 'Arial'
    
    # Add empty paragraph after banner for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(12)

def set_font_run(run, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33), bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def add_heading_styled(doc, text, level=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    if level == 2:
        run = p.add_run(text)
        set_font_run(run, name='Arial', size_pt=12.5, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    elif level == 3:
        run = p.add_run(text)
        set_font_run(run, name='Arial', size_pt=11, color_rgb=(0x0A, 0x84, 0x81), bold=True)
    return p

def add_paragraph_styled(doc, text="", space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        set_font_run(run, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33), italic=italic)
    return p

def add_bullet_point(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font_run(run, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))
    return p

def add_screenshot_box(doc, text_placeholder, caption):
    # Dotted single-cell table for screenshot
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(5.8)
    
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3FAF8")
    set_dashed_borders(table, "0A8481")
    set_cell_margins(cell, top=400, bottom=400, left=300, right=300)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Camera symbol / Icon
    icon_run = p.add_run("📷\n\n")
    set_font_run(icon_run, name='Segoe UI Symbol', size_pt=24, color_rgb=(0x0A, 0x84, 0x81))
    
    txt_run = p.add_run(text_placeholder)
    set_font_run(txt_run, name='Calibri', size_pt=11, color_rgb=(0x77, 0x77, 0x77), italic=True)
    
    # Add caption paragraph
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(6)
    p_cap.paragraph_format.space_after = Pt(18)
    cap_run = p_cap.add_run(caption)
    set_font_run(cap_run, name='Calibri', size_pt=10, color_rgb=(0x55, 0x55, 0x55), italic=True)

def build_report():
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69) # A4

    # -------------------------------------------------------------
    # PAGE 1: COVER
    # -------------------------------------------------------------
    p_top_spacer = doc.add_paragraph()
    p_top_spacer.paragraph_format.space_before = Pt(40)
    
    # Laporan Akhir Project Header
    p_laporan = doc.add_paragraph()
    p_laporan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_laporan.paragraph_format.space_after = Pt(4)
    run_laporan = p_laporan.add_run("LAPORAN AKHIR PROJECT")
    set_font_run(run_laporan, name='Arial', size_pt=13, color_rgb=(0x0A, 0x84, 0x81), bold=True)
    
    # Pemrograman Web II
    p_pw2 = doc.add_paragraph()
    p_pw2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_pw2.paragraph_format.space_after = Pt(12)
    run_pw2 = p_pw2.add_run("PEMROGRAMAN WEB II")
    set_font_run(run_pw2, name='Arial', size_pt=24, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    
    # Divider Line Table (Teal thin separator)
    div_table = doc.add_table(rows=1, cols=1)
    div_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    div_table.autofit = False
    div_table.columns[0].width = Inches(3.5)
    set_cell_shading(div_table.cell(0, 0), "0A8481")
    # Borderless
    tblPr = div_table._tbl.tblPr
    borders = parse_xml(f'<w:tblBorders {nsdecls("w")}><w:top w:val="none"/><w:bottom w:val="none"/><w:left w:val="none"/><w:right w:val="none"/></w:tblBorders>')
    tblPr.append(borders)
    # Minimise padding/height
    set_cell_margins(div_table.cell(0, 0), top=10, bottom=10, left=0, right=0)
    
    # App Title Spacer
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(28)
    
    # Judul Aplikasi Label
    p_judul_lbl = doc.add_paragraph()
    p_judul_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_judul_lbl.paragraph_format.space_after = Pt(4)
    run_judul_lbl = p_judul_lbl.add_run("Judul Aplikasi")
    set_font_run(run_judul_lbl, name='Calibri', size_pt=11, color_rgb=(0x66, 0x66, 0x66), italic=True)
    
    # App Name (MABIPRO)
    p_mabipro = doc.add_paragraph()
    p_mabipro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_mabipro.paragraph_format.space_after = Pt(36)
    run_mabipro = p_mabipro.add_run("MABIPRO — MANAJEMEN BISNIS PROPERTI")
    set_font_run(run_mabipro, name='Arial', size_pt=16, color_rgb=(0x0A, 0x84, 0x81), bold=True)
    
    # Logo Box
    logo_table = doc.add_table(rows=1, cols=1)
    logo_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    logo_table.autofit = False
    logo_table.columns[0].width = Inches(2.2)
    l_cell = logo_table.cell(0, 0)
    set_cell_shading(l_cell, "F9F9F9")
    set_dashed_borders(logo_table, "CCCCCC")
    set_cell_margins(l_cell, top=250, bottom=250, left=150, right=150)
    lp = l_cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    l_icon = lp.add_run("🏢\n")
    set_font_run(l_icon, name='Segoe UI Symbol', size_pt=24, color_rgb=(0x88, 0x88, 0x88))
    l_txt = lp.add_run("LOGO")
    set_font_run(l_txt, name='Arial', size_pt=10, color_rgb=(0x88, 0x88, 0x88), bold=True)
    
    # Author Info Spacer
    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_after = Pt(36)
    
    # Disusun Oleh
    p_oleh = doc.add_paragraph()
    p_oleh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_oleh.paragraph_format.space_after = Pt(4)
    run_oleh = p_oleh.add_run("Disusun oleh:")
    set_font_run(run_oleh, name='Calibri', size_pt=11, color_rgb=(0x66, 0x66, 0x66), italic=True)
    
    # Nama Lengkap
    p_nama = doc.add_paragraph()
    p_nama.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nama.paragraph_format.space_after = Pt(2)
    run_nama = p_nama.add_run("[ Nama Lengkap ]")
    set_font_run(run_nama, name='Arial', size_pt=12, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    
    # NIM
    p_nim = doc.add_paragraph()
    p_nim.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nim.paragraph_format.space_after = Pt(6)
    run_nim = p_nim.add_run("NIM. [ ....................... ]")
    set_font_run(run_nim, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))
    
    # Kelas & Dosen
    p_kelas_dosen = doc.add_paragraph()
    p_kelas_dosen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_kelas_dosen.paragraph_format.space_after = Pt(40)
    run_kd = p_kelas_dosen.add_run("Kelas: [ ...... ]   |   Dosen Pengampu: [ ................... ]")
    set_font_run(run_kd, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))
    
    # University Details
    p_prog = doc.add_paragraph()
    p_prog.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_prog.paragraph_format.space_after = Pt(3)
    run_prog = p_prog.add_run("[ PROGRAM STUDI ................................ ]")
    set_font_run(run_prog, name='Arial', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    
    p_fak = doc.add_paragraph()
    p_fak.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_fak.paragraph_format.space_after = Pt(3)
    run_fak = p_fak.add_run("[ FAKULTAS ................................ ]")
    set_font_run(run_fak, name='Arial', size_pt=11, color_rgb=(0x11, 0x1E, 0x38))
    
    p_univ = doc.add_paragraph()
    p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ.paragraph_format.space_after = Pt(4)
    run_univ = p_univ.add_run("[ NAMA UNIVERSITAS ]")
    set_font_run(run_univ, name='Arial', size_pt=11.5, color_rgb=(0x0A, 0x84, 0x81), bold=True)
    
    p_tahun = doc.add_paragraph()
    p_tahun.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tahun.paragraph_format.space_after = Pt(0)
    run_tahun = p_tahun.add_run("[ Tahun ]")
    set_font_run(run_tahun, name='Calibri', size_pt=11, color_rgb=(0x55, 0x55, 0x55))
    
    # End of Page 1
    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 2: IDENTITAS & DESKRIPSI PROJECT
    # -------------------------------------------------------------
    add_header_banner(doc, "I", "IDENTITAS & DESKRIPSI PROJECT")
    
    add_heading_styled(doc, "1.1 Identitas Project", level=2)
    
    # Table for Identitas Project
    table_id = doc.add_table(rows=7, cols=2)
    table_id.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_id.autofit = False
    table_id.columns[0].width = Inches(2.2)
    table_id.columns[1].width = Inches(4.3)
    set_table_borders(table_id, "CCCCCC")
    
    labels_id = [
        "Nama Aplikasi",
        "Nama Penyusun",
        "NIM",
        "Kelas",
        "Dosen Pengampu",
        "Link Repository GitHub\n(wajib bersifat public)",
        "Link Demo / Live URL"
    ]
    
    values_id = [
        "MABIPRO — Manajemen Bisnis Properti",
        "[ Nama Lengkap ]",
        "[ ......................................... ]",
        "[ ......................................... ]",
        "[ ......................................... ]",
        "[ https://github.com/username/nama-repo ]",
        "[ https://... (jika ada) ]"
    ]
    
    for r in range(7):
        # Format label cell
        cell_lbl = table_id.cell(r, 0)
        set_cell_margins(cell_lbl, top=100, bottom=100, left=100, right=100)
        p_lbl = cell_lbl.paragraphs[0]
        run_lbl = p_lbl.add_run(labels_id[r])
        set_font_run(run_lbl, name='Calibri', size_pt=10.5, color_rgb=(0x33, 0x33, 0x33), bold=True)
        
        # Format value cell
        cell_val = table_id.cell(r, 1)
        set_cell_margins(cell_val, top=100, bottom=100, left=100, right=100)
        p_val = cell_val.paragraphs[0]
        run_val = p_val.add_run(values_id[r])
        # Make links or placeholders a bit gray / italic
        is_placeholder = "[" in values_id[r] or "github.com" in values_id[r] or "https" in values_id[r]
        set_font_run(run_val, name='Calibri', size_pt=10.5, color_rgb=(0x66, 0x66, 0x66) if is_placeholder else (0x11, 0x1E, 0x38), italic=is_placeholder)

    # Note below the table
    p_note = doc.add_paragraph()
    p_note.paragraph_format.space_before = Pt(6)
    p_note.paragraph_format.space_after = Pt(14)
    note_run_bold = p_note.add_run("Catatan: ")
    set_font_run(note_run_bold, name='Calibri', size_pt=9.5, color_rgb=(0x0A, 0x84, 0x81), bold=True, italic=True)
    note_run_text = p_note.add_run("Pastikan repository GitHub diatur Public (Settings → General → Danger Zone → Change visibility) agar dapat diakses dan dinilai oleh dosen tanpa perlu undangan kolaborator.")
    set_font_run(note_run_text, name='Calibri', size_pt=9.5, color_rgb=(0x66, 0x66, 0x66), italic=True)
    
    add_heading_styled(doc, "1.2 Deskripsi Singkat Aplikasi", level=2)
    add_paragraph_styled(doc, "MABIPRO (Manajemen Bisnis Properti) adalah sistem manajemen bisnis properti berbasis web yang dirancang khusus untuk mempermudah pemantauan penjualan unit, progres pembangunan (konstruksi), berkas legalitas, dan laporan keuangan transaksi. Aplikasi ini menjembatani kolaborasi empat peran penting di perusahaan pengembang properti, yaitu Administrator, Sales/Marketing, Tim Produksi Lapangan, dan Tim Legalitas. Melalui integrasi data dalam satu sistem, perusahaan dapat menghindari kesalahan pencatatan data dan keterlambatan laporan fisik.")
    add_paragraph_styled(doc, "Sistem ini membantu memecahkan masalah administrasi properti konvensional yang sering kali tidak sinkron antara progres penjualan dari marketing dengan progres pembangunan di lapangan dari tim produksi. Dengan MABIPRO, status unit properti (Belum Terjual, Sudah DP, Terjual) diperbarui secara real-time. Informasi mengenai dokumen hukum per unit (SHM, IMB/PBG) juga langsung terpantau kelengkapannya. Aplikasi ini ditargetkan untuk pengembang perumahan skala kecil hingga menengah untuk mengotomatisasi alur bisnis harian.")
    
    add_heading_styled(doc, "1.3 Tujuan", level=2)
    add_bullet_point(doc, "Membangun sistem informasi terpusat untuk memantau siklus hidup unit perumahan dari awal hingga serah terima kunci.")
    add_bullet_point(doc, "Meningkatkan transparansi dan koordinasi antar-divisi (marketing, konstruksi, dan legalitas) secara real-time.")
    add_bullet_point(doc, "Menyediakan kalkulator simulasi angsuran KPR reaktif untuk membantu marketing memberikan estimasi keuangan yang akurat kepada calon pembeli.")
    add_bullet_point(doc, "Mempermudah dokumentasi progres konstruksi fisik dengan laporan riwayat tahapan pembangunan disertai foto bukti dari lapangan.")
    add_bullet_point(doc, "Mengarsipkan berkas legalitas penting (SHM, IMB, dll.) per unit secara digital agar aman dan mudah diakses.")
    
    # End of Page 2
    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 3: TEKNOLOGI YANG DIGUNAKAN
    # -------------------------------------------------------------
    add_header_banner(doc, "II", "TEKNOLOGI YANG DIGUNAKAN")
    
    add_paragraph_styled(doc, "Berikut daftar teknologi, framework, dan tools yang digunakan dalam pengembangan aplikasi MABIPRO:")
    
    # Table for Technologies
    table_tech = doc.add_table(rows=8, cols=3)
    table_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tech.autofit = False
    table_tech.columns[0].width = Inches(1.8)
    table_tech.columns[1].width = Inches(2.7)
    table_tech.columns[2].width = Inches(2.0)
    set_table_borders(table_tech, "CCCCCC")
    
    # Table Tech Headers
    headers_tech = ["Kategori", "Teknologi / Tools", "Keterangan / Versi"]
    for c in range(3):
        cell_h = table_tech.cell(0, c)
        set_cell_shading(cell_h, "0A8481")
        set_cell_margins(cell_h, top=120, bottom=120, left=100, right=100)
        p_h = cell_h.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_h = p_h.add_run(headers_tech[c])
        set_font_run(run_h, name='Arial', size_pt=10, color_rgb=(0xFF, 0xFF, 0xFF), bold=True)
        
    tech_data = [
        ("Bahasa Pemrograman", "PHP, JavaScript", "PHP v8.3, ES6+"),
        ("Framework / Library", "Laravel Framework, Livewire", "Laravel v11.x, Livewire v3.x"),
        ("Frontend", "Tailwind CSS, Alpine.js, Blade View", "Tailwind v4.0.0, Alpine v3.15+"),
        ("Database", "MySQL Database Server", "MySQL v8.0"),
        ("Web Server / Runtime", "Node.js (Vite Bundle Compiler)", "Node v18+, Vite v8.0.0"),
        ("Tools Pendukung", "Git CLI, Visual Studio Code, Composer, npm", "Version Control & Packages"),
        ("Hosting / Deployment", "Localhost (Laragon Development Env)", "Windows Web Server Suite")
    ]
    
    for r in range(1, 8):
        row_data = tech_data[r-1]
        for c in range(3):
            cell_c = table_tech.cell(r, c)
            set_cell_margins(cell_c, top=100, bottom=100, left=100, right=100)
            p_c = cell_c.paragraphs[0]
            run_c = p_c.add_run(row_data[c])
            set_font_run(run_c, name='Calibri', size_pt=10.5, color_rgb=(0x33, 0x33, 0x33))
            
    # Empty paragraph after table
    p_tech_space = doc.add_paragraph()
    p_tech_space.paragraph_format.space_before = Pt(8)
    
    add_heading_styled(doc, "2.1 Alasan Pemilihan Teknologi", level=2)
    
    add_paragraph_styled(doc, "Pemilihan teknologi di atas didasarkan pada efisiensi pengembangan aplikasi dan skalabilitas sistem:")
    
    p_lar = add_paragraph_styled(doc)
    run_lar_title = p_lar.add_run("1. Laravel 11 (PHP): ")
    set_font_run(run_lar_title, name='Calibri', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    run_lar_desc = p_lar.add_run("Menyediakan pondasi backend MVC (Model-View-Controller) yang kokoh, sistem routing yang aman, otentikasi siap pakai, serta ORM Eloquent yang memudahkan transaksi data relasional yang aman antara tabel Blok, Unit, Legalitas, dan Foto Progress.")
    set_font_run(run_lar_desc, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))
    
    p_lw = add_paragraph_styled(doc)
    run_lw_title = p_lw.add_run("2. Livewire: ")
    set_font_run(run_lw_title, name='Calibri', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    run_lw_desc = p_lw.add_run("Memungkinkan interaksi frontend yang dinamis dan reaktif (Single Page Application-like) langsung menggunakan kode PHP tanpa perlu menulis REST API atau menginstal framework Javascript kompleks (seperti React/Vue). Contohnya terlihat pada modal kalkulator simulasi KPR yang reaktif saat marketing menginput data nominal.")
    set_font_run(run_lw_desc, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))
    
    p_tw = add_paragraph_styled(doc)
    run_tw_title = p_tw.add_run("3. Tailwind CSS & Vite: ")
    set_font_run(run_tw_title, name='Calibri', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    run_tw_desc = p_tw.add_run("Tailwind CSS digunakan untuk merancang antarmuka (UI) admin dan dasbor secara responsif dan premium dengan utilitas kelas CSS yang ringkas. Vite mempercepat proses kompilasi aset frontend selama masa pengembangan.")
    set_font_run(run_tw_desc, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))
    
    p_db = add_paragraph_styled(doc)
    run_db_title = p_db.add_run("4. MySQL: ")
    set_font_run(run_db_title, name='Calibri', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    run_db_desc = p_db.add_run("Database relasional open-source paling populer, andal untuk menyimpan data transaksional penjualan rumah, menghitung unit tersisa, dan mengelola dokumen legalitas.")
    set_font_run(run_db_desc, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))

    # End of Page 3
    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 4 & 5: FITUR APLIKASI
    # -------------------------------------------------------------
    add_header_banner(doc, "III", "FITUR APLIKASI")
    
    add_paragraph_styled(doc, "Bagian ini menampilkan fitur-fitur utama aplikasi MABIPRO beserta bukti berupa screenshot pendukung:")
    
    # Fitur 1
    add_heading_styled(doc, "Fitur 1: Dashboard Multi-Role (Admin, Marketing, Produksi, Legalitas)", level=3)
    add_paragraph_styled(doc, "Halaman beranda utama yang menyajikan statistik dinamis sesuai peran (role) pengguna. Admin mendapatkan visualisasi grafik penjualan unit secara menyeluruh, sedangkan tim produksi, marketing, dan legalitas melihat ringkasan tugas spesifik yang harus diselesaikan di dasbor mereka.")
    add_screenshot_box(doc, "Tempel screenshot dashboard di sini", "Gambar 1. Halaman Dashboard Utama Multi-Role MABIPRO")
    
    # Fitur 2
    add_heading_styled(doc, "Fitur 2: Manajemen Penjualan & Simulasi Pembayaran (Cash & KPR)", level=3)
    add_paragraph_styled(doc, "Modul yang digunakan marketing untuk mengubah status penjualan unit (Belum Terjual, Sudah DP, Terjual) dan melakukan simulasi pembayaran. Dilengkapi dengan input interaktif untuk menghitung otomatis nominal DP, total pokok kredit, dan rincian angsuran bulanan KPR.")
    add_screenshot_box(doc, "Tempel screenshot form modul pembayaran & simulasi KPR di sini", "Gambar 2. Form Simulasi Pembayaran Unit Properti (KPR & Cash)")
    
    # Page Break before page 5
    doc.add_page_break()
    
    # Fitur 3
    add_heading_styled(doc, "Fitur 3: Pemantauan Progres Konstruksi & Riwayat Foto Lapangan", level=3)
    add_paragraph_styled(doc, "Modul bagi divisi produksi untuk melacak persentase progres pembangunan unit rumah (0-100%). Petugas lapangan dapat mengunggah bukti foto nyata dari lapangan serta memberikan catatan untuk tiap tahapan konstruksi (Pondasi, Struktur, Finishing, Serah Terima).")
    add_screenshot_box(doc, "Tempel screenshot riwayat dan update progres konstruksi di sini", "Gambar 3. Fitur Riwayat Tahapan dan Unggah Foto Progres Konstruksi")
    
    # Fitur 4
    add_heading_styled(doc, "Fitur 4: Manajemen Dokumen Legalitas Unit", level=3)
    add_paragraph_styled(doc, "Modul bagi divisi legalitas untuk memverifikasi dokumen hukum penting untuk setiap unit perumahan. Petugas dapat mengunggah, mengunduh, dan melihat pratinjau (preview) dokumen seperti SHM, IMB, PBB, dan Akta Jual Beli. Status kelengkapan dokumen diperbarui secara otomatis.")
    add_screenshot_box(doc, "Tempel screenshot kelola dokumen legalitas di sini", "Gambar 4. Halaman Unggah dan Verifikasi Berkas Dokumen Legalitas")

    # End of Page 4/5
    doc.add_page_break()

    # -------------------------------------------------------------
    # PAGE 6: KESIMPULAN
    # -------------------------------------------------------------
    add_header_banner(doc, "IV", "KESIMPULAN")
    
    add_heading_styled(doc, "4.1 Kesimpulan", level=2)
    add_paragraph_styled(doc, "Aplikasi MABIPRO (Manajemen Bisnis Properti) telah berhasil dirancang dan dibangun dengan mengintegrasikan empat divisi operasional penting pengembang properti. Seluruh modul utama telah berfungsi dengan baik dan lolos uji fungsionalitas: dasbor multi-role yang aman, modul manajemen penjualan & simulasi KPR interaktif (Marketing), pengelolaan dokumen legalitas (Legalitas), serta pemantauan riwayat konstruksi bermedia foto lapangan (Produksi).")
    add_paragraph_styled(doc, "Dengan menggunakan Laravel 11 dan Livewire, sistem dapat menyajikan antarmuka yang responsif, interaktif, dan mudah dirawat. Implementasi MABIPRO ini berhasil memecahkan kendala administrasi properti konvensional dengan menyelaraskan data fisik progres pembangunan dan administrasi berkas secara real-time, sehingga membantu manajemen developer mengambil keputusan bisnis yang lebih cepat dan akurat.")
    
    add_heading_styled(doc, "4.2 Kendala & Saran Pengembangan", level=2)
    
    p_ken = add_paragraph_styled(doc)
    run_ken_title = p_ken.add_run("Kendala Pengerjaan: ")
    set_font_run(run_ken_title, name='Calibri', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    run_ken_desc = p_ken.add_run("Penanganan penyimpanan berkas foto progres dan PDF dokumen legalitas berukuran besar memerlukan optimasi memori dan ruang penyimpanan server (seperti kompresi gambar otomatis sebelum diunggah). Sinkronisasi status penjualan unit antar-dashboard divisi yang aktif juga masih mengandalkan reload/polling database.")
    set_font_run(run_ken_desc, name='Calibri', size_pt=11, color_rgb=(0x33, 0x33, 0x33))
    
    p_sar = add_paragraph_styled(doc)
    run_sar_title = p_sar.add_run("Saran Pengembangan ke Depan:\n")
    set_font_run(run_sar_title, name='Calibri', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True)
    
    add_bullet_point(doc, "Notifikasi WhatsApp & Email: Menambahkan pengiriman pesan otomatis kepada konsumen ketika berkas legalitas selesai diverifikasi atau saat progres konstruksi unit mereka mencapai tahap baru.")
    add_bullet_point(doc, "Integrasi Payment Gateway: Menyediakan modul pembayaran uang muka (DP) online menggunakan Midtrans agar pencatatan di marketing terverifikasi secara instan.")
    add_bullet_point(doc, "Ekspor Laporan Keuangan & Konstruksi: Menambahkan tombol ekspor seluruh data unit, progress, dan pembayaran dalam format lembar kerja Excel (.xlsx) untuk mempermudah audit perusahaan.")
    
    # Extra Space
    p_lamp_space = doc.add_paragraph()
    p_lamp_space.paragraph_format.space_before = Pt(12)
    
    # Lampiran Section
    p_lamp = doc.add_paragraph()
    p_lamp.paragraph_format.space_after = Pt(4)
    run_lamp_bold = p_lamp.add_run("Lampiran (opsional):")
    set_font_run(run_lamp_bold, name='Calibri', size_pt=11, color_rgb=(0x11, 0x1E, 0x38), bold=True, italic=True)
    
    add_bullet_point(doc, "Link Video Demo: [ Tulis Link Video Demo di Sini, misal: YouTube ]")
    add_bullet_point(doc, "Source Code: [ Tulis Link GitHub Repository Proyek di Sini ]")
    add_bullet_point(doc, "Dokumentasi Tambahan: [ Tulis Link Google Drive / Dropbox jika ada ]")

    # Save document
    filename = "Laporan_Akhir_MABIPRO.docx"
    doc.save(filename)
    print(f"Success! {filename} generated successfully.")

if __name__ == "__main__":
    build_report()
